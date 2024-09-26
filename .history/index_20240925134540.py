import pandas as pd
from docx import Document

# Read the Excel file with names
excel_file = 'names_list.xlsx'  # Replace with your Excel file path
df = pd.read_excel(excel_file)

# Open the DOCX template
template_file = 'offer_template.docx'  # Replace with your DOCX template path
template = Document(template_file)

# Path where the new offer letters will be saved
output_folder = './offer_letters/'  # Ensure this folder exists or create it

# Loop over each name in the Excel file
for index, row in df.iterrows():
    name = row['Name']
    
    # Create a copy of the document
    doc = Document(template_file)
    
    # Replace the placeholder in the document with the actual name
    for paragraph in doc.paragraphs:
        if '{Name}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{Name}', name)
    
    # Save the new offer letter with the recipient's name in the filename
    output_path = f"{output_folder}Offer_Letter_{name}.docx"
    doc.save(output_path)

    print(f"Offer letter for {name} saved as {output_path}")

print("Offer letters created for all names.")
